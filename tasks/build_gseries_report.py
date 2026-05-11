"""Build the G-Series presentation .docx for the BI-RADS mammography project.

All numbers are taken verbatim from:
- tasks/cascade_log.md
- tasks/lessons.md (Lesson #49, #50)
- outputs/cascade/evaluation_report.md / evaluation_metrics.json
- experiments/EXPERIMENTS.md
- CLAUDE.md (dataset statistics)
"""

from pathlib import Path

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Pt, RGBColor

OUT_PATH = Path("/home/alilivan.turk/Desktop/YeniYaklaşım/tasks/G_Series_Sunum.docx")

# -----------------------------------------------------------------------------
# Styling helpers
# -----------------------------------------------------------------------------

NAVY = RGBColor(0x0B, 0x33, 0x66)
RED = RGBColor(0xB0, 0x1C, 0x1C)
GREEN = RGBColor(0x14, 0x6B, 0x3A)
GRAY = RGBColor(0x55, 0x55, 0x55)


def set_cell_bg(cell, color_hex):
    from docx.oxml.ns import nsdecls
    from docx.oxml import parse_xml

    tc_pr = cell._tc.get_or_add_tcPr()
    shd = parse_xml(rf'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>')
    tc_pr.append(shd)


def add_heading(doc, text, level=1, color=None):
    h = doc.add_heading(text, level=level)
    if color is not None:
        for r in h.runs:
            r.font.color.rgb = color
    return h


def add_para(doc, text, *, bold=False, italic=False, size=11, color=None, align=None):
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    r = p.add_run(text)
    r.bold = bold
    r.italic = italic
    r.font.size = Pt(size)
    if color is not None:
        r.font.color.rgb = color
    return p


def add_bullets(doc, items, *, bold_first_word=False):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        if bold_first_word and ":" in item:
            head, rest = item.split(":", 1)
            r1 = p.add_run(head + ":")
            r1.bold = True
            p.add_run(rest)
        else:
            p.add_run(item)


def add_table(doc, header, rows, *, header_color="0B3366", first_col_bold=False,
              col_widths=None, highlight_rows=None):
    """highlight_rows: dict {row_idx: hex_color} (0-based, body rows only)."""
    table = doc.add_table(rows=1 + len(rows), cols=len(header))
    table.style = "Light Grid Accent 1"
    table.autofit = True

    hdr = table.rows[0].cells
    for i, h in enumerate(header):
        hdr[i].text = ""
        p = hdr[i].paragraphs[0]
        r = p.add_run(h)
        r.bold = True
        r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        r.font.size = Pt(10)
        set_cell_bg(hdr[i], header_color)
        hdr[i].vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    for ri, row in enumerate(rows):
        cells = table.rows[ri + 1].cells
        for ci, val in enumerate(row):
            cells[ci].text = ""
            p = cells[ci].paragraphs[0]
            r = p.add_run(str(val))
            r.font.size = Pt(10)
            if first_col_bold and ci == 0:
                r.bold = True
            cells[ci].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        if highlight_rows and ri in highlight_rows:
            for c in cells:
                set_cell_bg(c, highlight_rows[ri])

    if col_widths is not None:
        for col_idx, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[col_idx].width = w

    doc.add_paragraph()  # spacing
    return table


def page_break(doc):
    doc.add_page_break()


# -----------------------------------------------------------------------------
# Document build
# -----------------------------------------------------------------------------

doc = Document()

# Page margins
for section in doc.sections:
    section.left_margin = Cm(2.0)
    section.right_margin = Cm(2.0)
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(2.0)

# =============================================================================
# COVER
# =============================================================================

title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = title.add_run("G-SERİSİ KASKAD DENEYLERİ")
r.bold = True
r.font.size = Pt(28)
r.font.color.rgb = NAVY

sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = sub.add_run("Mammografi BI-RADS Sınıflandırması")
r.font.size = Pt(16)
r.font.color.rgb = GRAY

sub2 = doc.add_paragraph()
sub2.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = sub2.add_run("Üç Aşamalı Yumuşak Kaskad — Deney Raporu ve Bulgular")
r.italic = True
r.font.size = Pt(13)
r.font.color.rgb = GRAY

doc.add_paragraph()
doc.add_paragraph()

meta = doc.add_paragraph()
meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = meta.add_run("Tarih: 2026-04-29\nProje: Multi-View Hierarchical Mammography Classifier\n"
                 "Şampiyon Model: C6 (test macro F1 = 0.6762)\n"
                 "G-Serisi Sonuç: REDDEDİLDİ (-4.96pp)")
r.font.size = Pt(12)

doc.add_paragraph()

box = doc.add_paragraph()
box.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = box.add_run(
    "Bu rapor, BI-RADS 1/2/4/5 sınıflandırması için tasarlanan üç aşamalı yumuşak "
    "kaskad mimarisinin (G1 + G2a + G2b) C6 baseline'ına karşı yapılan kapsamlı "
    "değerlendirmesini içerir. Tüm sayısal değerler MLflow run'larından, "
    "outputs/cascade/ artifaktlarından ve project lessons (Lesson #49, #50) "
    "dosyalarından doğrudan alınmıştır."
)
r.italic = True
r.font.size = Pt(10)
r.font.color.rgb = GRAY

page_break(doc)

# =============================================================================
# 1. YÖNETİCİ ÖZETİ
# =============================================================================

add_heading(doc, "1. Yönetici Özeti", level=1, color=NAVY)

add_para(
    doc,
    "G-Serisi, C6 şampiyon modelinin (tek-omurga + multi-task hiyerarşi) ötesine "
    "geçmek için tasarlanan ilk yapısal müdahaledir. Üç ayrı uzman model — bir ikili "
    "(benign/malign) yönlendirici ve iki adet sınıf-içi uzman (BR1↔BR2 ve BR4↔BR5) — "
    "C6'nın train/val ayrımının alt kümeleri üzerinde bağımsız olarak eğitilmiş, "
    "ardından inference sırasında olasılıklar P(BRk) = P(stage1=parent) × P(BRk | stage1=parent) "
    "şeklinde çarpımsal olarak birleştirilmiştir.",
)

add_para(doc, "Özet bulgular:", bold=True)

add_bullets(
    doc,
    [
        "Validation gate'leri geçti: G1 = 0.9664 (eşik 0.93), G2a = 0.7020, G2b = 0.7740. Hepsi C6'nın ilgili alt-küme metriklerinin üzerinde.",
        "Test sonucu C6'ya göre düştü: yumuşak kaskad macro F1 = 0.6266 (Δ = -0.0496 / -4.96pp).",
        "BR2 sınıfında çöküş (-12.2pp) tek başına macro F1'in büyük kısmını yedi: BR2→BR1 drift'i %13'ten %28'e çıktı.",
        "Yumuşak kaskad ≈ Sert kaskad (Δ +0.0004pp): tüm aşamalar >%95 güvenle çıkış üretti, çarpımsal birleşim argmax × argmax'a çöktü.",
        "Asıl başarısızlık nedeni: train→test prior shift'in üç aşama boyunca çarpımsal olarak büyümesi (Lesson #50).",
    ],
)

add_para(doc, "Karar:", bold=True)
add_bullets(
    doc,
    [
        "C6 şampiyon kalır, kaskad deploy edilmez.",
        "G2a/G2b LR-schedule fix'i ile retrain edilmeyecek (yapısal sorun, eğitim kararsızlığı değil).",
        "Bir sonraki deney: F2 — logit-adjusted training (Menon ve diğerleri, 2021), prior shift'i doğrudan hedefler.",
    ],
)

page_break(doc)

# =============================================================================
# 2. PROJE BAĞLAMI VE VERİ KÜMESİ
# =============================================================================

add_heading(doc, "2. Proje Bağlamı ve Veri Kümesi", level=1, color=NAVY)

add_heading(doc, "2.1 Görev Tanımı", level=2)
add_para(
    doc,
    "Hasta bazlı 4-sınıf BI-RADS sınıflandırması (1/2/4/5; BI-RADS-3 yoktur). "
    "Girdi: 1 hasta = 4 görüntü (RCC, LCC, RMLO, LMLO). Görüntüler 1024×1024, "
    "8-bit grayscale PNG, DICOM → segmentasyon → windowing → tight crop → CLAHE → "
    "letterbox boru hattıyla üretilir.",
)

add_heading(doc, "2.2 Veri Dağılımı", level=2)

add_table(
    doc,
    ["Bölüm", "Veri Kümesi", "Hasta", "Görüntü"],
    [
        ["Train", "Dataset_1024_8bit (85%)", "8.557 → 7.273 train", "34.228"],
        ["Val", "Dataset_1024_8bit (15%, seed=42)", "1.284", "—"],
        ["Test", "Dataset_Test_1024_8bit (sabit)", "1.655", "6.620"],
    ],
    first_col_bold=True,
)

add_para(
    doc,
    "Train sınıf dağılımı (hasta sayıları): BR1 = 1.678, BR2 = 2.754, BR4 = 1.898, BR5 = 2.227. "
    "Test sınıf dağılımı (hasta sayıları): BR1 = 163, BR2 = 596, BR4 = 288, BR5 = 608.",
)

add_para(
    doc,
    "KRİTİK GÖZLEM — train ve test prior'ları aynı değildir:",
    bold=True,
    color=RED,
)
add_bullets(
    doc,
    [
        "Train BR1 oranı %19.6 → Test %9.9 (yarıya düşüyor).",
        "Train BR2 oranı %32.2 → Test %36.0 (artıyor).",
        "Train BR4 oranı %22.2 → Test %17.4 (azalıyor).",
        "Train BR5 oranı %26.0 → Test %36.7 (artıyor).",
        "Bu prior shift, tüm `sqrt-inv-frequency` class weight kalibrasyonlarını test setinde sistematik olarak yanlış kılar.",
    ],
)

add_heading(doc, "2.3 Kullanılan Train İstatistikleri (8-bit)", level=2)
add_para(
    doc,
    "Tüm-piksel mean/std = 0.1210 / 0.1977 (train ve test ~aynı: 0.1210 vs 0.1237). "
    "Doku (nonzero) mean/std = 0.3512 / 0.1804. Sınıf ağırlıkları (sqrt-inv): "
    "[1.28, 1.00, 1.20, 1.11] → [BR1, BR2, BR4, BR5].",
)

page_break(doc)

# =============================================================================
# 3. ÖNCEKİ DENEYLER VE C6'NIN DOĞUŞU
# =============================================================================

add_heading(doc, "3. Şampiyon C6'ya Giden Yol", level=1, color=NAVY)

add_para(
    doc,
    "G-Serisi'nden önce A → B → C → D ablation serileri çalıştırılmıştır. "
    "Aşağıdaki tablo, EXPERIMENTS.md indeksinden alınan en kritik 8 deneyi içerir.",
)

add_table(
    doc,
    ["Deney", "Tarih", "Val F1", "Test F1", "Backbone", "Ana Değişiklik"],
    [
        ["a1 (focal)", "2026-04-09", "0.7108", "0.6270", "ConvNeXtV2-L", "8-bit + focal loss"],
        ["a1_ce", "2026-04-09", "0.7141", "0.6370", "ConvNeXtV2-L", "Focal → CE (Lesson #28)"],
        ["b1", "2026-04-09", "0.7334", "0.6387", "ConvNeXtV2-L", "Hiyerarşik head + 4-class CE"],
        ["b5", "2026-04-09", "0.7286", "0.6615", "ConvNeXtV2-L", "+ SWA (start_epoch=5)"],
        ["c1", "2026-04-12", "0.7158", "0.6431", "ConvNeXtV2-L", "Loss ağırlık ayarı"],
        ["c4 (base)", "2026-04-12", "0.7178", "0.6269", "ConvNeXtV2-Base", "Kapasite düşürme (Lesson #25)"],
        ["c6 ★ ŞAMPİYON", "2026-04-12", "0.7183", "0.6762", "ConvNeXtV2-L", "Asymmetry loss → 0 (Lesson #22)"],
        ["c7", "2026-04-12", "0.7261", "0.6468", "ConvNeXtV2-L", "C6'nın varyantı, val↑ test↓"],
    ],
    first_col_bold=True,
    highlight_rows={6: "DDEEDD"},
)

add_para(doc, "C6'yı şampiyon yapan kararlar (cumulative):", bold=True)
add_bullets(
    doc,
    [
        "Loss tipi: CrossEntropy (focal loss bu veri setinde zararlı — Lesson #28, #37).",
        "Hiyerarşik 4 head: binary + benign-sub + malign-sub + full, ağırlıklar 0.10 / 0.45 / 0.45.",
        "SWA: start_epoch=5, AveragedModel + update_bn (+1.47pp katkı — Lesson #37).",
        "Asymmetry loss: KAPALI (weight=0.0) — bilateral fusion zaten F_diff token'ı içeriyor (Lesson #22).",
        "Backbone: ConvNeXtV2-Large (Base versiyonu test setinde -4.93pp regresse oluyor — Lesson #25).",
        "Sınıf ağırlıkları: sqrt-inv frequency, max=1.0 normalize (Lesson #27).",
    ],
)

add_para(
    doc,
    "D-serisi (d1–d7), C6'nın her bir bileşenini ayrı ayrı kapatarak negatif kontroller "
    "üretti: hiyerarşinin her parçasının çıkarılması test F1'i düşürdü (Lesson #37). "
    "Bu, C6'nın \"goldilocks\" konfigürasyonu olduğunu doğrular.",
)

page_break(doc)

# =============================================================================
# 4. G-SERİSİ HİPOTEZ VE MİMARİ
# =============================================================================

add_heading(doc, "4. G-Serisi Hipotezi ve Mimarisi", level=1, color=NAVY)

add_heading(doc, "4.1 Motivasyon", level=2)
add_para(
    doc,
    "C6 sonrası analizler şu darboğazları işaret etti:",
)
add_bullets(
    doc,
    [
        "C6'nın binary head'i val 0.9530 / test 0.939 — güçlü ama tek omurga 4 görevi paylaştığı için BR1↔BR2 sınırı yumuşak kalıyor.",
        "BR4↔BR5 sınırı, paylaşımlı backbone tarafından düşük ağırlıkla öğreniliyor (BR4→BR5 drift'i %36.1 — Lesson #44).",
        "Hipotez: Her sınıf çiftine kendi backbone'unu adamak, paylaşımlı omurgada sıkışan ayırt edici özellikleri açığa çıkarır.",
    ],
)

add_heading(doc, "4.2 Mimari (Üç Aşamalı Yumuşak Kaskad)", level=2)

# ASCII-style architecture description (rendered as text)
arch_para = doc.add_paragraph()
arch_para.paragraph_format.left_indent = Cm(0.5)
r = arch_para.add_run(
    "                     ┌── G2a: ConvNeXtV2-L + lateral + bilateral fusion\n"
    "                     │       + 2-class head (BR1 vs BR2)\n"
    "                     │       → P(BR1|benign), P(BR2|benign)\n"
    "  G1 (Stage-1) ──────┤\n"
    "  ConvNeXtV2-L       │\n"
    "  + mean pool        │\n"
    "  + 2-class head     └── G2b: ConvNeXtV2-L + lateral + bilateral fusion\n"
    "  → P(benign),               + 2-class head (BR4 vs BR5)\n"
    "    P(malign)                → P(BR4|malign), P(BR5|malign)\n"
    "\n"
    "  Yumuşak birleşim:  P(BRk) = P(stage1 = parent) × P(BRk | parent)\n"
)
r.font.name = "Consolas"
r.font.size = Pt(9)

add_heading(doc, "4.3 Aşama Tasarımları", level=2)

add_table(
    doc,
    ["Aşama", "Eğitim Verisi", "Mimari", "Çıkış"],
    [
        ["G1 (Stage-1)", "7.273 train / 1.284 val (tüm hastalar, ikili etiket)",
         "Backbone → mean-pool (4 görüntü × S spatial token) → 2-class head. FÜZYON YOK (claude.md doktrini).",
         "P(benign), P(malign)"],
        ["G2a (Stage-2a)", "3.767 train / 665 val (yalnız BR1/BR2)",
         "C6'nın tüm füzyon yığını + 2-class head", "P(BR1|benign), P(BR2|benign)"],
        ["G2b (Stage-2b)", "3.506 train / 619 val (yalnız BR4/BR5)",
         "C6'nın tüm füzyon yığını + 2-class head", "P(BR4|malign), P(BR5|malign)"],
    ],
    first_col_bold=True,
)

add_heading(doc, "4.4 Doğrulanmış Tasarım Kararları (Livan ile mutabakat)", level=2)
add_bullets(
    doc,
    [
        "Stage-1 sıfırdan eğitildi (C6'dan fine-tune değil) — temiz ablasyon.",
        "Üç aşamada da ConvNeXtV2-Large kullanıldı (Lesson #25: küçük backbone bu rejimde zararlı).",
        "Augmentation C6 ile birebir aynı — kaskad ayrışması izole edilebilsin diye.",
        "SWA üç aşamada da açık (swa_start_epoch=5).",
        "Yumuşak kaskad birincil; sert kaskad sağlık-kontrolü olarak rapor edildi.",
        "Class weight'ler her aşamanın kendi train alt-dağılımına göre sqrt-inv hesaplandı: G1=[1.04, 1.00], G2a=[1.28, 1.00], G2b=[1.08, 1.00].",
    ],
)

page_break(doc)

# =============================================================================
# 5. EĞİTİM SÜREÇLERİ VE PHASE E (VAL GATE'LERİ)
# =============================================================================

add_heading(doc, "5. Eğitim Süreci ve Validation Gate'leri", level=1, color=NAVY)

add_heading(doc, "5.1 Phase E — Validation Sonuçları", level=2)

add_para(doc, "Sanity gate'leri (eğitim sonrası, test öncesi):", italic=True)

add_table(
    doc,
    ["Aşama", "Best Val F1 Macro", "Eşik", "C6 Referansı", "Sonuç"],
    [
        ["G1 (binary)", "0.9664", "≥ 0.93", "C6 binary 0.9530", "GEÇTİ (+1.34pp vs C6)"],
        ["G2a (BR1↔BR2)", "0.7020", "≥ C6 alt-küme tahmini", "~0.71 (cross-FP düzeltmeli)", "MARGINAL — sınırda"],
        ["G2b (BR4↔BR5)", "0.7740", "≥ C6 alt-küme tahmini", "~0.755 (cross-FP düzeltmeli)", "GEÇTİ (+1.9pp)"],
    ],
    first_col_bold=True,
    highlight_rows={0: "DDEEDD", 2: "DDEEDD"},
)

add_para(doc, "Val seviyesinde elde edilen pozitif sinyaller:", bold=True, color=GREEN)
add_bullets(
    doc,
    [
        "Malign sınır (BR4↔BR5): G2b val F1 = 0.7740, C6'nın 4-class macro alt sınırı 0.7413'ün ÜZERİNDE. Uzman özellikleri burada işe yaradı.",
        "İkili sınır (benign↔malign): G1 val F1 = 0.9664, C6 binary head 0.9530'dan +1.34pp daha iyi — füzyonsuz tasarım val'de daha iyi öğrendi.",
        "G2a benign sınırı (BR1↔BR2): 0.7020 — sadece eşiğin sınırında, kaskad hipotezinin bu eksende henüz desteklenmediğini gösteriyor.",
    ],
)

add_heading(doc, "5.2 Eğitim Patolojisi — OneCycleLR Yanlış Konfigürasyonu", level=2)

add_para(
    doc,
    "G2a ve G2b'de bir \"yükseliş-ve-çöküş\" trajektörisi gözlendi: peak val'den "
    "sonra her epoch çoğunluk sınıfını tahmin etmeye geri döndü.",
)

add_table(
    doc,
    ["Aşama", "Trivial Faz", "Yükseliş", "Peak", "Çöküş"],
    [
        ["G2a", "ep 1-5 (0.383)", "ep 6-13 (0.55→0.70)", "ep 13 (0.7020)", "ep 14-33 (0.383)"],
        ["G2b", "ep 1-3 (0.35)", "ep 4-21 (0.71-0.77 osc.)", "ep 23 (0.7740)", "ep 24-43 (0.35)"],
    ],
    first_col_bold=True,
)

add_para(doc, "Kök sebep:", bold=True)
add_bullets(
    doc,
    [
        "C6'nın LR schedule'ı kopyalandı: epochs=100, max_lr=5e-4, pct_start=0.3 → backbone effective max_lr 1e-4 epoch 30'da.",
        "C6 patience=20 ile epoch 26'da early-stop oldu, peak LR'ye HİÇ ulaşmadı.",
        "G2a peak val epoch 13'te → erken durdurma sayacı başladı → epoch 33'e kadar koştu (peak LR'ye ulaştı, model çöktü).",
        "G2b benzer: epoch 23 peak → epoch 43 çöküş.",
        "SWA bile yardım etmedi: ortalama post-collapse epoch'lardan alındığı için final_using_swa=False.",
    ],
)

add_para(
    doc,
    "Bu, kayda alınmış bir öğrenim olarak Lesson #49'a eklendi: \"Bir LR schedule'ı "
    "deneyler arasında kopyalarken, hedef deneyin beklenen early-stopping epoch'una "
    "göre yeniden ölçeklendirilmelidir.\" Genel kural: peak fazı beklenen durma "
    "noktasının hemen ÖTESİNDE olacak şekilde, epochs ≈ 3.3 × e* seçilmelidir.",
    italic=True,
    color=GRAY,
)

page_break(doc)

# =============================================================================
# 6. PHASE F + G — TEST SONUÇLARI
# =============================================================================

add_heading(doc, "6. Phase F + G — Test Sonuçları (Olumsuz)", level=1, color=NAVY)

add_heading(doc, "6.1 Headline Metrikler", level=2)

add_table(
    doc,
    ["Metrik", "C6 (Şampiyon)", "G-Serisi Yumuşak", "G-Serisi Sert", "Δ vs C6"],
    [
        ["Test macro F1", "0.6762", "0.6266", "0.6262", "-0.0496"],
        ["Test weighted F1", "—", "0.6826", "0.6821", "—"],
        ["Test binary F1 (implied)", "0.939", "0.9309", "—", "-0.0081"],
    ],
    first_col_bold=True,
    highlight_rows={0: "F8DDDD"},
)

add_para(doc, "Yumuşak kaskad ≈ Sert kaskad (Δ +0.0004pp):", bold=True)
add_para(
    doc,
    "Tüm aşamalar >%95 güvenle çıkış üretiyor. Yumuşak çarpım argmax×argmax'a "
    "çöküyor. Yumuşak kompozisyonun sağladığı varsayılan belirsizlik harmanlaması "
    "bu rejimde mevcut değil — uzmanlar yeterince \"yumuşak belirsiz\" değil.",
)

add_heading(doc, "6.2 Sınıf Bazlı F1 Karşılaştırması", level=2)

add_table(
    doc,
    ["Sınıf", "Support", "C6 F1", "Kaskad F1", "Δ", "C6 P/R", "Kaskad P/R"],
    [
        ["BR1", "163", "0.531", "0.496", "-3.5pp", "0.517 / 0.546", "0.389 / 0.681"],
        ["BR2", "596", "0.798", "0.676", "-12.2pp", "0.829 / 0.770", "0.877 / 0.550"],
        ["BR4", "288", "0.518", "0.517", "-0.1pp", "0.515 / 0.521", "0.448 / 0.611"],
        ["BR5", "608", "0.857", "0.818", "-3.9pp", "0.837 / 0.878", "0.821 / 0.814"],
    ],
    first_col_bold=True,
    highlight_rows={1: "F8DDDD"},  # BR2 catastrophic
)

add_para(doc, "BR2'deki -12.2pp tek başına macro F1'in 3pp+ kaybını açıklıyor.", bold=True, color=RED)

add_heading(doc, "6.3 Test Confusion Matrix (Yumuşak Kaskad)", level=2)

add_table(
    doc,
    ["", "pred BR1", "pred BR2", "pred BR4", "pred BR5"],
    [
        ["true BR1 (n=163)", "111", "39", "12", "1"],
        ["true BR2 (n=596)", "166", "328", "94", "8"],
        ["true BR4 (n=288)", "8", "5", "176", "99"],
        ["true BR5 (n=608)", "0", "2", "111", "495"],
    ],
    first_col_bold=True,
)

add_para(doc, "Kritik drift desenleri:", bold=True)
add_bullets(
    doc,
    [
        "BR2 → BR1 = %27.85 (kaskad) vs %13 (C6) — yaklaşık 2× artış. BR2'nin precision'ı korunuyor (0.877) ama recall %55'e düşüyor: 166 BR2 hastası BR1'e gönderildi.",
        "BR1 over-prediction: 285 tahmin / 163 gerçek = 1.75× over-prediction. Recall %68.1'e çıkıyor ama precision 0.389'a çakılıyor — sahte zafer.",
        "BR4 → BR5 = %34.4 (kaskad) ≈ %36.1 (C6) — değişiklik yok, malign morfoloji boundary halen kara kutu.",
    ],
)

page_break(doc)

# =============================================================================
# 7. KÖK SEBEP ANALİZİ
# =============================================================================

add_heading(doc, "7. Neden Başarısız Oldu? — Üç Yapısal Sebep", level=1, color=NAVY)

add_heading(doc, "7.1 Sebep 1: Train→Test Prior Shift'in Çarpımsal Büyümesi", level=2)
add_para(
    doc,
    "Bu, Lessons #47, #48 ve #50'de tekrar tekrar saptanan dataset'in baskın "
    "genelleme vergisidir. Kaskad bunu üç katına çıkardı:",
)
add_bullets(
    doc,
    [
        "Stage 1 (binary): train benign:malign = 1.07, test = 0.85 → G1'in class weight'i [1.04, 1.00] hafifçe yanlış kalibre.",
        "Stage 2a (BR1↔BR2): train BR1:BR2 = 0.61, test = 0.27 → G2a'nın BR1 class weight 1.28 SİSTEMATİK OLARAK ÇOK AGRESİF.",
        "Stage 2b (BR4↔BR5): train BR4:BR5 = 0.85, test = 0.47 → G2b'nin BR4 class weight 1.08 yine prior'a göre çok yüksek.",
        "Üç yanlış kalibre edilmiş aşamanın çarpımı, tek bir 4-class head'in tek-aşamalı hatasından çok daha büyüktür.",
    ],
)

add_heading(doc, "7.2 Sebep 2: Bileşke (Compositional) Hata Yüzeyinin Genişlemesi", level=2)
add_para(
    doc,
    "C6'da bir BR2→BR1 hatası, full head'in BR1 logit'inin BR2'yi geçmesini "
    "gerektirir. Kaskadda aynı hata İKİ farklı yoldan oluşabilir:",
)
add_bullets(
    doc,
    [
        "Yol A: G1 doğru ('benign') + G2a yanlış ('BR1 > BR2'). Bu sıkça oluyor (BR2'nin %55 recall).",
        "Yol B: G1 yanlış ('malign') + sonra herhangi bir route. Daha nadir ama sıfır değil.",
        "Disjunctive ('VEYA') hata yüzeyi, conjunctive ('VE') hata yüzeyinden geniş — ki tek başlı C6'nın kullandığı budur.",
    ],
)

add_heading(doc, "7.3 Sebep 3: Uzman Özellikleri Test'te Transfer Etmedi", level=2)
add_para(
    doc,
    "Lesson #50'nin en şaşırtıcı bulgusu: G2b'nin malign axisindeki +3.3pp val "
    "kazanımı test'te BR4 için statistikİ olarak sıfıra geriledi (0.517 vs 0.518), "
    "BR5'te -3.9pp kayıp oldu.",
)
add_bullets(
    doc,
    [
        "Val→test gap'i C6'da 4.55pp idi. G2b'de bu daha da büyüdü.",
        "G1'in füzyonsuz mimarisi val'de +1.34pp kazandırdı, test'te -0.81pp kaybettirdi (Lesson #50). Füzyon bir REGÜLARİZER görevi görüyor — sadece bir özellik kanalı değil.",
        "Sonuç (Lesson #48 ile uyumlu): Auxiliary head'lerin değeri eğitim-zamanı mimarisel bir regularizasyondur, post-hoc kompozisyonel bir özellik değil. Onları kaldırıp uzmanlaştırmak, hierarchical reconstruction'ın inference-time versiyonuyla aynı hatadır.",
    ],
)

page_break(doc)

# =============================================================================
# 8. NE İŞE YARADI / NE İŞE YARAMADI
# =============================================================================

add_heading(doc, "8. Pozitif ve Negatif Etkilerin Toplu Tablosu", level=1, color=NAVY)

add_heading(doc, "8.1 POZİTİF Etkiler", level=2, color=GREEN)

add_table(
    doc,
    ["Bileşen / Karar", "Gözlemlenen Etki", "Kanıt"],
    [
        ["Manifest tabanlı veri ayrımı", "Cascade train/val, C6'nın seed=42 split'inin tutarlı bir alt-kümesi oldu — apple-to-apple karşılaştırma sağladı.",
         "tools/build_cascade_manifests.py + sanity assertion'lar"],
        ["G1'in füzyonsuz tasarımı (val)", "C6 binary head'i +1.34pp geçti (0.9530 → 0.9664). Multi-task gradient kompetisyonu kaldırılınca binary öğrenme keskinleşti.",
         "MLflow run ad4526d7… val_f1_macro"],
        ["G2b uzmanlaşması (val)", "BR4↔BR5 ekseninde C6'nın 4-class lower bound'unu +3.3pp aştı.",
         "MLflow run c5926b3c… val_f1_macro = 0.7740"],
        ["Wrapper-model yaklaşımı", "C6 kod yoluna 0 satır müdahale edildi. Reusable: MultiViewBackbone, BilateralLateralFusion, BilateralFusion vs.",
         "models/cascade_model.py (~150 LOC)"],
        ["SWA G1'de", "G1 stable trained, using_swa=True. Backbone-only kıyasla füzyon yok = LR daha az destabilize etti.",
         "G1 final checkpoint metadata"],
        ["Phase E gate disiplini", "Test inference'a geçmeden val gate'leri ile kalibrasyon hatasını yakalayabilirdik (G2a'nın marginal sonucu erken sinyaldi).",
         "tasks/cascade_log.md Phase E"],
    ],
    first_col_bold=True,
)

add_heading(doc, "8.2 NEGATİF Etkiler", level=2, color=RED)

add_table(
    doc,
    ["Bileşen / Karar", "Gözlemlenen Etki", "Kanıt"],
    [
        ["Üç-aşamalı yumuşak kaskad", "Test macro F1 -4.96pp. Tek başına yapısal başarısızlık.",
         "evaluation_metrics.json"],
        ["Class weight çarpımı", "Üç aşamanın train-prior'a kalibre weight'leri test prior shift altında çarpımsal hata üretti.",
         "Lesson #50 root cause analysis"],
        ["LR schedule kopyalama (G2a/G2b)", "epochs=100 + pct_start=0.3, early-stopping ile peak LR fazına ulaşan G2a/G2b'yi çökertti. SWA averaging zarar verdi.",
         "Lesson #49"],
        ["BR2 sınıfında çöküş", "BR2 F1 -12.2pp (0.798 → 0.676). BR2→BR1 drift %13'ten %28'e iki katına çıktı.",
         "outputs/cascade/evaluation_report.md"],
        ["G1'in füzyonsuz mimarisi (test)", "Val'de kazanan, test'te -0.81pp kaybeden — füzyon REGÜLARİZER işi görüyormuş; kaldırılması overfit'i artırdı.",
         "Lesson #50: val→test gap C6'da 1.4pp, G1'de 3.55pp"],
        ["Yumuşak ≈ Sert kaskad", "Soft composition'ın belirsizlik harmanlama avantajı (+0.0004pp) sıfır. Aşamalar aşırı güvenli (%95+).",
         "evaluation_metrics.json hard vs soft"],
        ["G2a'nın tek-epoch peak'i", "Val 0.7020 yalnızca tek epoch'tan geliyor; trajektori kararsız. Final checkpoint kırılgan.",
         "Lesson #49 LR collapse trajectories"],
        ["Compute maliyeti (~2.5× C6)", "Üç ayrı backbone eğitimi yaklaşık 60+ GPU-saat, 0 test kazanımı.",
         "tasks/cascade_log.md compute budget"],
    ],
    first_col_bold=True,
)

add_heading(doc, "8.3 Önceki Deneylerden Devralınan KISITLAR (uygulamada doğrulandı)", level=2)
add_bullets(
    doc,
    [
        "Asymmetry loss (Lesson #22): Kalıcı olarak kapalı, kaskad da kullanmadı.",
        "Backbone küçültme (Lesson #25): Reddedildi, üç aşama da Large kullanıldı.",
        "Class weight manipülasyonu (Lesson #27): Sıfır-toplam kuralı kaskad genelinde de geçerli olduğu kanıtlandı.",
        "Focal loss (Lesson #28, #37): Kaskad CE kullandı.",
        "Mixup (Lesson #34): Kaskad da kullanmadı.",
    ],
)

page_break(doc)

# =============================================================================
# 9. ÖĞRENİMLER (LESSON #49 + #50 ÖZETİ)
# =============================================================================

add_heading(doc, "9. Çıkarılan Öğrenimler — Lesson #49 ve #50", level=1, color=NAVY)

add_heading(doc, "9.1 Lesson #49 — Eğitim-zamanı Bulgular", level=2)
add_bullets(
    doc,
    [
        "Auxiliary-head kaldırma, özellik-zengin sınırlarda (G2b/malign morfoloji) kullanılabilir bir uzman üretir; özellik-fakir sınırlarda (G2a/BR1↔BR2) ayırt edici sinyali yaratamaz.",
        "Bir LR schedule'ı yeni bir deneye kopyarken, beklenen early-stopping epoch'una göre yeniden boyutlandırılmalıdır. epochs ≈ 3.3 × e* genel kuralı.",
        "final_using_swa = False, eğitim kararsızlığının lider göstergesidir — SWA averaging regime change'ler arasında yapıldıysa val-best epoch'tan kötü çıkar.",
        "Cross-class FP analizi, 4-class baseline'ı binary subset metriğine doğru kalibre etmenin doğru yoludur.",
    ],
)

add_heading(doc, "9.2 Lesson #50 — Test-zamanı Bulgular ve Karar", level=2)
add_bullets(
    doc,
    [
        "Bu veri setinde kaskad ayrıştırma, 0.6762'nin ÖTESİNE bir yol değildir. Kaskad araştırma yönü kapatıldı.",
        "Mimarisel basitleştirmeden (G1 no-fusion) gelen val kazanımları, val→test gap analizi yapılmadan kabul edilmemelidir. +1.34pp val + 2.15pp gap genişlemesi = NET NEGATİF.",
        "Train→test prior shift bu veri setinin baskın genelleme vergisidir (Tasks 1.3, 1.4, G-series tümü buna takıldı).",
        "Yumuşak kaskad bileşkesi bu rejimde belirsizlik-harmanlama avantajı sağlamaz; aşamalar >%95 güvenli olduğu sürece soft = hard.",
    ],
)

add_heading(doc, "9.3 Önceki Lesson'larla Bağlantılar", level=2)
add_bullets(
    doc,
    [
        "Lesson #22 (asymmetry loss off): Doğrulandı.",
        "Lesson #25 (Large backbone şart): Doğrulandı.",
        "Lesson #27 (class weight zero-sum): Kaskad aşamalar arası genelleştirildi.",
        "Lesson #37 (C6 goldilocks): Bir negatif sonuç daha eklendi.",
        "Lesson #44 (C6 baseline): Aynen doğrulandı.",
        "Lesson #47 (val-tuned offsets fail under prior shift): Daha dramatik biçimde doğrulandı.",
        "Lesson #48 (auxiliary heads' value training-time): Yapısal versiyonda da doğrulandı.",
    ],
)

page_break(doc)

# =============================================================================
# 10. SONRAKİ ADIM: F2
# =============================================================================

add_heading(doc, "10. Sonraki Adım — F2: Logit-Adjusted Training", level=1, color=NAVY)

add_para(
    doc,
    "G-Serisi'nin reddedilmesi sonrası, root cause analizi (train→test prior shift) "
    "doğrudan hedef alacak bir deney scaffold'lanmıştır.",
)

add_heading(doc, "10.1 F2 Yöntemi", level=2)
add_bullets(
    doc,
    [
        "Menon ve diğerleri (2021) — Logit Adjustment: eğitim-zamanı logit'lere `tau × log pi_train` shift uygula, inference'ta raw logit ile tahmin et.",
        "pi_train = [1678, 2754, 1898, 2227]/8557 → log pi = [-1.6293, -1.1336, -1.5060, -1.3460].",
        "Test etiketleri ASLA incelenmez (no Saerens-style estimation, no test-prior injection). Fairness-clean.",
        "Üç tau değeri sweep ediliyor: 0.5, 1.0, 1.5. Val'da en iyi tau seçilip tek test raporu üretilecek.",
    ],
)

add_heading(doc, "10.2 Hedefler ve Kabul Kriterleri", level=2)
add_bullets(
    doc,
    [
        "Test macro F1 > 0.6762 (C6'yı geçmek) → F2 KAZANIR, Lesson #51 olarak belgelenir.",
        "BR2 recall ≥ 0.78 (C6'da 0.74).",
        "Phase E gate: val_full_f1_macro ≥ 0.6783 her tau için.",
        "Phase F: Test labels'a bakmadan tau* val ile seçilir; tie-break son 5 epoch val F1 std'sine göre.",
    ],
)

add_heading(doc, "10.3 Codebase Değişiklikleri (zaten scaffold'lı)", level=2)
add_bullets(
    doc,
    [
        "utils/logit_adjustment.py — yeni dosya: train priors + helpers.",
        "utils/losses.py — MultiHeadLoss `logit_adjustment` dict'i kabul ediyor; full head'e shift uygulanıyor; LA açıkken full head'in class_weights düşürülüyor (Menon recipe).",
        "configs/experiment_v2_birads/convnextv2_large_8bit_F2_la_tau{05,10,15}.yaml — üç config.",
        "scripts/slurm/train_F2_la_tau{05,10,15}.sh — SLURM job templates.",
        "BLAST RADIUS: 0 — train.py, models/, data/, utils/metrics.py değişmedi.",
    ],
)

page_break(doc)

# =============================================================================
# 11. ARTIFAKTLAR VE DOĞRULAMA İZİ
# =============================================================================

add_heading(doc, "11. Artifaktlar ve Doğrulama İzi", level=1, color=NAVY)

add_heading(doc, "11.1 MLflow Run'ları (DagsHub)", level=2)

add_table(
    doc,
    ["Run", "Run ID", "Experiment"],
    [
        ["G1 (Stage-1 binary)", "ad4526d7ef684e7e845ea977aa49d4a2", "cascade/stage1_binary"],
        ["G2a (Stage-2a BR1↔BR2)", "bae58239a2c34f81b76c4f14a5cbbe04", "cascade/stage2_benign"],
        ["G2b (Stage-2b BR4↔BR5)", "c5926b3c3e8841faaae1892678ce97ca", "cascade/stage2_malign"],
        ["C6 referans", "6859aed2a37e43b8b72b5333b2573275", "birads-1024-8bit-ablation"],
    ],
    first_col_bold=True,
)

add_heading(doc, "11.2 Disk Artifaktları", level=2)
add_bullets(
    doc,
    [
        "outputs/cascade/test_probs.parquet — bileşke test olasılıkları (1655 hasta × 4 sınıf).",
        "outputs/cascade/evaluation_report.md — Phase G insan-okunur raporu.",
        "outputs/cascade/evaluation_metrics.json — tüm metrikler (yumuşak + sert).",
        "checkpoints/cascade/{G1,G2a,G2b}_best.pt — sabit alias'lar.",
        "data/manifests/cascade/*.csv — Phase A manifest'leri (sanity assertion'lı).",
        "models/cascade_model.py — wrapper class CascadeStageModel (~150 LOC).",
        "train_cascade.py — bağımsız trainer (train.py'a dokunmadı).",
        "tools/cascade_inference.py + cascade_evaluate.py — Phase F + G script'leri.",
    ],
)

add_heading(doc, "11.3 Kod Değişikliği Disiplini", level=2)
add_para(
    doc,
    "G-Serisi tamamen ek (additive) bir eksen olarak implement edildi. C6 kod yoluna "
    "(train.py, models/full_model.py, models/classification_heads.py) tek satır "
    "değişiklik yapılmadı. Bu, kaskad reddedildiği halde C6'nın reproduce "
    "edilebilirliğini koruduğu anlamına gelir — ileride paper write-up için kritik.",
    italic=True,
    color=GRAY,
)

page_break(doc)

# =============================================================================
# 12. KAPANIŞ
# =============================================================================

add_heading(doc, "12. Kapanış ve Sunum Notları", level=1, color=NAVY)

add_heading(doc, "12.1 Hocalara Önemli Mesajlar", level=2)
add_bullets(
    doc,
    [
        "Negatif sonuç paper-grade bir bulgu: \"3-aşamalı yumuşak kaskad test ettik; aşama-bazında val iyileşmeleri (G1, G2b), bileşke val→test prior shift ve G2a'daki kırılgan tek-epoch peak nedeniyle test'e transfer etmedi. Paylaşımlı-omurga multi-task tasarım (C6) bu BI-RADS veri setinde en güçlü tek-model konfigürasyonu olmaya devam ediyor.\"",
        "Çalışmanın yöntemsel disiplini güçlü: pre-registered Phase E gate'leri, train.py'a dokunmayan wrapper kod, full MLflow + W&B logging, manifest tabanlı reproducibility.",
        "Bir sonraki deney (F2) doğrudan tespit edilen kök sebebi (prior shift) hedef alıyor — kaskadın başarısızlığından öğrendiklerimizden besleniyor.",
    ],
)

add_heading(doc, "12.2 Açık Sorular", level=2)
add_bullets(
    doc,
    [
        "Saerens-style test-time prior correction (val-estimated prior) bir alternatif olabilir mi? F2 sonuçlarına bağlı.",
        "16-bit pipeline (Tier-2 Task 2.1) F2'den sonra mı? 300GB regen + 17h GPU run; F2 0.70'i geçmezse devreye girer.",
        "BR1↔BR2 sınırı bir feature-engineering problemi mi yoksa veri toplama problemi mi? Lesson #27 + #49 ikincisini düşündürtüyor.",
        "Calibration improvement (temperature scaling şu an gradyan almıyor — Lesson #44) bağımsız bir kazanım kaynağı olabilir mi?",
    ],
)

doc.add_paragraph()
doc.add_paragraph()

footer = doc.add_paragraph()
footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = footer.add_run(
    "── Rapor Sonu ──\n\n"
    "Hazırlayan: Senior AI Research Scientist (Claude, Sonnet/Opus harness)\n"
    "Veri kaynakları: tasks/cascade_log.md, tasks/lessons.md (Lesson #49, #50), "
    "outputs/cascade/, experiments/EXPERIMENTS.md, CLAUDE.md\n"
    "Tarih: 2026-04-29"
)
r.italic = True
r.font.size = Pt(9)
r.font.color.rgb = GRAY

# =============================================================================
# Save
# =============================================================================

doc.save(str(OUT_PATH))
print(f"Saved: {OUT_PATH}")
print(f"Size: {OUT_PATH.stat().st_size:,} bytes")
